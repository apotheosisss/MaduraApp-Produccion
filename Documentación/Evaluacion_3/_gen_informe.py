# -*- coding: utf-8 -*-
"""Genera el Informe Técnico Eval 3 (.docx) de MaduraApp."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = RGBColor(0x2E, 0x7D, 0x32)
DARK = RGBColor(0x1A, 0x1C, 0x19)

doc = Document()

# Estilos base
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(hdr[i], "2E7D32")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(8.5)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def para(text, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    return p

# ─────────────── PORTADA ───────────────
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MaduraApp"); r.bold = True; r.font.size = Pt(40); r.font.color.rgb = GREEN
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sistema de Análisis de Madurez Agrícola mediante Visión Computacional")
r.font.size = Pt(14); r.italic = True
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Informe Técnico — Estado de Avance 3"); r.bold = True; r.font.size = Pt(18)
for _ in range(6):
    doc.add_paragraph()
info = [
    ("Asignatura", "Taller Aplicado de Programación (TPY1101)"),
    ("Sección", "001D"),
    ("Estudiante", "Claudio Vicente Aro Kath — RUT 22.022.498-8"),
    ("Docente guía", "José Ignacio Campos Arévalo"),
    ("Fecha", "Junio 2026"),
    ("Repositorio", "github.com/apotheosisss/MaduraApp-Produccion"),
]
ti = doc.add_table(rows=0, cols=2); ti.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in info:
    c = ti.add_row().cells
    c[0].text = ""; rr = c[0].paragraphs[0].add_run(k); rr.bold = True
    c[1].text = v
doc.add_page_break()

# ─────────────── ÍNDICE (campo TOC) ───────────────
h1("Índice")
par = doc.add_paragraph()
run = par.add_run()
fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'TOC \\o "1-2" \\h \\z \\u'
fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'separate')
t_run = OxmlElement('w:t'); t_run.text = "Actualiza este índice en Word con clic derecho → Actualizar campos."
fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'end')
for e in (fld1, instr, fld2, t_run, fld3):
    run._r.append(e)
doc.add_page_break()

# ─────────────── 1. INTRODUCCIÓN ───────────────
h1("1. Introducción")
para("Las pérdidas post-cosecha en frutas climatéricas (aguacate, plátano, tomate, mango) "
     "representan en Chile entre un 20% y un 40% de la producción según FAO/ODEPA. Una causa "
     "raíz es la falta de criterios objetivos y accesibles para determinar el punto óptimo de "
     "consumo o procesamiento de la fruta.")
para("MaduraApp es un sistema de visión computacional accesible desde el móvil que clasifica el "
     "estado de madurez de cuatro frutas climatéricas en tres categorías (Inmaduro / Óptimo / "
     "Sobre maduro) y entrega recomendaciones agronómicas en tiempo real. Este informe "
     "corresponde al Estado de Avance 3, cuyo foco es el aseguramiento de calidad: el plan de "
     "pruebas, su aplicación a los componentes del proyecto y las mejoras derivadas de los "
     "resultados.")

# ─────────────── 2. RESUMEN EVAL 1 ───────────────
h1("2. Resumen Evaluación 1 — Problema y requisitos")
for b in [
    "Problemática: pérdidas post-cosecha por falta de criterios objetivos de madurez (Ishikawa).",
    "Solución: app móvil + IA que clasifica madurez y recomienda acción.",
    "Requisitos funcionales (RF-01 a RF-14): captura/galería, selección de fruta, inferencia, "
    "recomendación, persistencia, historial paginado y offline, semáforo visual, validación de "
    "formato, health check, autenticación JWT y feedback.",
    "KPI principal: precisión del modelo mAP@50 ≥ 0.75.",
]:
    doc.add_paragraph(b, style="List Bullet")

# ─────────────── 3. RESUMEN EVAL 2 ───────────────
h1("3. Resumen Evaluación 2 — Arquitectura y producto")
for b in [
    "Stack: Android nativo (Kotlin + CameraX + Room + MVVM); FastAPI (Python 3.12 async); "
    "YOLO26n; SQLAlchemy + Alembic; PostgreSQL/SQLite.",
    "Arquitectura 4+1 (Kruchten): vistas lógica, de procesos, de desarrollo, física y de escenarios.",
    "Producto: selector de fruta, cámara y galería, semáforo de madurez, historial offline; "
    "backend con /v1/predict, /v1/history, /v1/health.",
    "Modelo YOLO26n entrenado 80 épocas: mAP@50 = 0.9229, 5.2 MB.",
]:
    doc.add_paragraph(b, style="List Bullet")
para("Sobre esa base, la Evaluación 3 incorporó autenticación JWT, feedback con rating, "
     "endurecimiento de seguridad (OWASP) y un rediseño de UI (Material 3), todos sometidos a pruebas.")

# ─────────────── 4. EVAL 3 ───────────────
h1("4. Evaluación 3 — Pruebas y mejoras")

h2("4.1 Plan de pruebas")
para("Se confeccionó un plan de pruebas alineado a la problemática, con 36 casos automatizados "
     "(17 backend + 19 Android) clasificados en validación, verificación, seguridad y operacionales.")

para("Pruebas de validación — Autenticación (backend)")
add_table(
    ["ID", "Funcionalidad", "Acción / datos", "Esperado", "Obtenido", "Estado"],
    [
        ["CP-01", "Registro rechaza email duplicado", "POST /auth/register dup", "HTTP 409", "409", "✅"],
        ["CP-02", "Login rechaza password incorrecta", "POST /auth/login", "HTTP 401", "401", "✅"],
        ["CP-03", "Registro exitoso emite JWT", "POST /auth/register válido", "201 + token", "201+token", "✅"],
    ],
)
para("Pruebas de validación — Inferencia (backend)")
add_table(
    ["ID", "Funcionalidad", "Acción / datos", "Esperado", "Obtenido", "Estado"],
    [
        ["CP-04", "Health check público", "GET /health", "200, model_loaded", "200", "✅"],
        ["CP-05", "Predicción con fruta detectada", "POST /predict (JWT)", "200, success, scan_id", "200+id", "✅"],
        ["CP-06", "Predicción con filtro de fruta", "POST /predict mango", "200, success", "200", "✅"],
        ["CP-07", "Imagen sin fruta detectable", "POST /predict", "200, success:false", "success:false", "✅"],
        ["CP-08", "Formato no soportado", "POST /predict .gif", "400", "400", "✅"],
        ["CP-09", "Tipo de fruta inválido", "POST /predict manzana", "400", "400", "✅"],
    ],
)
para("Pruebas de validación — Historial y Feedback (backend)")
add_table(
    ["ID", "Funcionalidad", "Acción / datos", "Esperado", "Obtenido", "Estado"],
    [
        ["CP-10", "Historial vacío usuario nuevo", "GET /history", "200, []", "[]", "✅"],
        ["CP-11", "Historial refleja escaneo", "GET /history", "contiene registro", "OK", "✅"],
        ["CP-12", "Paginación limit/offset", "GET /history?limit=5", "200, ≤5", "OK", "✅"],
        ["CP-13", "Validación límite máximo", "GET /history?limit=200", "422", "422", "✅"],
        ["CP-14", "Registrar rating 1-5", "POST /feedback", "201", "201", "✅"],
        ["CP-15", "Rechazar rating inválido", "POST /feedback", "422", "422", "✅"],
    ],
)
para("Pruebas de seguridad (OWASP)")
add_table(
    ["ID", "Funcionalidad", "OWASP", "Esperado", "Obtenido", "Estado"],
    [
        ["CP-16", "/predict exige JWT", "A01", "401 sin token", "401", "✅"],
        ["CP-17", "/history exige JWT", "A01", "401 sin token", "401", "✅"],
        ["CP-18", "Política de contraseña robusta", "A07", "422 a débil", "422", "✅"],
    ],
)
para("Pruebas de validación — Android (capa de datos y presentación): 19 pruebas JVM "
     "(FruitRepository: cache offline, refresh, health; ScanViewModel y HistoryViewModel: "
     "estados Idle/Loading/Success/NoDetection/Error). Todas en verde (CP-19 a CP-37).")

para("Pruebas de verificación (atributos de calidad)")
add_table(
    ["ID", "Atributo", "KPI", "Obtenido", "Estado"],
    [
        ["VER-01", "Precisión (mAP@50)", "≥ 0.75", "0.9229", "✅"],
        ["VER-02", "Recall por clase", "≥ 0.65", "> 0.80", "✅"],
        ["VER-03", "Tamaño del modelo", "< 10 MB", "5.2 MB", "✅"],
        ["VER-04", "Latencia inferencia CPU", "< 400 ms p95", "160-300 ms", "✅"],
        ["VER-05", "CI ejecuta la suite", "verde", "backend_ci.yml", "✅"],
        ["VER-06", "Sin secretos en repo", "0", "JWT por entorno", "✅"],
    ],
)
para("Pruebas operacionales: el backend aún no está desplegado en el AWS Laboratory del docente "
     "(pendiente). Las pruebas operacionales se documentan contra el entorno local (uvicorn + adb "
     "reverse / túnel) y se replicarán contra la URL de AWS una vez desplegado.", italic=True)

h2("4.2 Base de datos de pruebas")
para("Las pruebas usan una base efímera y aislada: SQLite in-memory para el backend, con "
     "fixtures que crean un usuario de prueba, un JWT real e imágenes JPEG sintéticas; el modelo "
     "YOLO se reemplaza por un mock. En Android se usan dobles de prueba (MockK) de la API y de "
     "Room. Ningún dato real ni secreto se utiliza en las pruebas; las contraseñas de prueba se "
     "hashean con bcrypt igual que en producción.")

h2("4.3 Aplicación de pruebas y resultados")
para("Ejecutadas el 21/06/2026: backend 17/17 (pytest, 5.34 s) y Android 19/19 "
     "(gradlew testDebugUnitTest). Verificaciones de calidad: mAP@50 = 0.9229, modelo 5.2 MB, "
     "APK compila, migraciones aplican. Total: 36/36 pruebas en verde.")
add_table(
    ["Suite", "Framework", "Pruebas", "Resultado"],
    [
        ["Backend (FastAPI)", "pytest 9.0.3", "17", "✅ 17 passed"],
        ["Android FruitRepository", "JUnit + MockK", "9", "✅ 9 passed"],
        ["Android ScanViewModel", "JUnit + coroutines-test", "6", "✅ 6 passed"],
        ["Android HistoryViewModel", "JUnit + LiveData", "4", "✅ 4 passed"],
        ["TOTAL", "", "36", "✅ 36 passed, 0 failed"],
    ],
)

h2("4.4 Mejoras al producto")
para("Los hallazgos de las pruebas y de una auditoría OWASP originaron 15 mejoras trazables a "
     "commits, mapeadas a los estándares de calidad de la industria.")
add_table(
    ["#", "Hallazgo", "Mejora aplicada", "Estándar", "Verificación"],
    [
        ["M-01", "JWT_SECRET_KEY hardcodeado (A02/A05)", "Secreto fuerte por entorno; rechaza arrancar con default en prod", "Seguridad", "Test manual prod"],
        ["M-02", "CORS '*' permisivo (A05)", "CORS acotado; métodos/headers explícitos", "Seguridad", "Config valida"],
        ["M-03", "Contraseña débil (A07)", "Mínimo 8 + letra + número (cliente y servidor)", "Seguridad", "CP-18 ✅"],
        ["M-04", "Tráfico HTTP en claro (A02)", "HTTPS forzado (network_security_config)", "Seguridad", "APK compila"],
        ["M-05", "Token sin cifrar en dispositivo (A02/M9)", "EncryptedSharedPreferences AES-256", "Seguridad", "19/19 ✅"],
        ["M-06", "Usuario no percibe la protección", "Indicador 'datos cifrados' honesto en login/registro", "Usabilidad/Seguridad", "Revisión visual"],
        ["M-07", "Emojis como íconos (UX)", "Íconos vectoriales + color M3 completo", "Usabilidad", "Render consistente"],
        ["M-08", "Modo oscuro inconsistente", "Modo oscuro real (values-night) + tokens", "Usabilidad", "Contraste OK"],
        ["M-09", "Sin feedback táctil; tipografía irregular", "MaterialCardView + roles M3 + 4/8dp", "Usabilidad", "19/19 ✅"],
        ["M-10", "Tests rotos tras añadir JWT", "Fixture auth_headers; códigos 401", "Corrección", "17/17 ✅"],
        ["M-11", "Build Android roto tras merge", "Código muerto eliminado; firmas alineadas", "Corrección", "APK + 19/19 ✅"],
        ["M-12", "Dependencias auth ausentes/incompat.", "pydantic[email]; bcrypt==4.0.1", "Corrección/Completitud", "17/17 ✅"],
        ["M-13", "Auth/feedback sin integrar", "Integración end-to-end por usuario", "Completitud", "CP-01..18 ✅"],
        ["M-14", "Recomendaciones del dominio", "Mensajes por estado y fruta; filtro mejora precisión", "Pertinencia", "CP-05/06 ✅"],
        ["M-15", "Sin integración continua", "backend_ci.yml ejecuta pytest en cada push", "Corrección/Completitud", "GitHub Actions"],
    ],
)

h2("4.5 Control de versiones")
para("Git + GitHub, 40 commits, Conventional Commits, ramas de feature, CI (backend_ci.yml) que "
     "corre la suite en cada push, y copias de configuración versionadas sin secretos "
     "(.env.example, requirements.txt, migraciones Alembic, build.gradle.kts, .gitignore).")

# ─────────────── 5. CONCLUSIÓN ───────────────
h1("5. Conclusión y lecciones aprendidas")
para("MaduraApp es hoy un producto funcional, probado y endurecido en seguridad. Las 36 pruebas "
     "pasan al 100%, las mejoras derivadas de los resultados cubren los cinco estándares de "
     "calidad, y la protección de datos personales (cifrado en tránsito/reposo, hashing, JWT) es "
     "real y verificable.")
para("Lecciones aprendidas:")
for b in [
    "Integrar temprano evita deuda: la rama de feature acumuló conflictos por separarse de main.",
    "Las pruebas son una red de seguridad real: detectaron las regresiones al añadir JWT.",
    "La seguridad se diseña con un marco (OWASP), no se improvisa.",
    "Honestidad técnica: se descartó etiquetar 'extremo a extremo' por ser falso; se usó 'cifrado en tránsito'.",
    "Fijar versiones de dependencias (bcrypt 4.0.1) evita fallos no reproducibles.",
    "El diseño se apoya en un sistema de tokens (Material 3), no en pantallas sueltas.",
    "La documentación se construyó sobre hechos verificables (tests y commits reales).",
]:
    doc.add_paragraph(b, style="List Bullet")
para("Pendiente honesto: desplegar el backend en el AWS Laboratory del docente, automatizar "
     "pruebas E2E y obtener la aprobación formal del plan de pruebas en la defensa.")

import os
out = os.path.join(os.path.dirname(__file__), "Informe_Tecnico_Evaluacion_3_MaduraApp.docx")
doc.save(out)
print("Guardado:", out)
